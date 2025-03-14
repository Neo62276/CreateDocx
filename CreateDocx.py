import pandas as pd
from docx import Document
import os
import logging
import re
import tkinter as tk
from tkinter import filedialog, messagebox

# 配置日志记录
logging.basicConfig(level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s')

# 函数：替换段落中的占位符
def replace_paragraph_placeholder(paragraph, placeholder_dict):
    try:
        # 构建正则表达式模式
        pattern = re.compile("|".join(re.escape(placeholder) for placeholder in placeholder_dict.keys()))
        # 合并段落中所有 run 的文本
        full_text = ''.join([run.text for run in paragraph.runs])
        # 进行替换
        new_text = pattern.sub(lambda m: placeholder_dict[m.group(0)], full_text)
        # 清空所有 run 的文本
        for run in paragraph.runs:
            run.text = ''
        # 将替换后的文本重新分配到第一个 run 中
        if paragraph.runs:
            paragraph.runs[0].text = new_text
    except Exception as e:
        logging.error(f"替换段落 {paragraph.text[:20]}... 中占位符时出错: {e}")

# 函数：替换表格中的占位符
def replace_table_placeholder(table, placeholder_dict):
    from docx.shared import Pt  # 导入 Pt 类用于设置字体大小
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for row in table.rows:
        # 确保行的 XML 元素有 w:trPr 节点
        trPr = row._element.xpath('.//w:trPr')
        if not trPr:
            trPr = OxmlElement('w:trPr')
            row._element.append(trPr)
        else:
            trPr = trPr[0]
        # 检查是否已有 w:cantSplit 节点，没有则添加
        cant_split = trPr.xpath('.//w:cantSplit')
        if not cant_split:
            cant_split = OxmlElement('w:cantSplit')
            trPr.append(cant_split)

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                # 合并段落中所有 run 的文本
                full_text = ''.join([run.text for run in paragraph.runs])
                # 构建正则表达式模式
                pattern = re.compile("|".join(re.escape(placeholder) for placeholder in placeholder_dict.keys()))
                # 进行替换
                new_text = pattern.sub(lambda m: placeholder_dict[m.group(0)], full_text)
                # 处理分隔符 ; 并拆分内容
                parts = new_text.split(';')
                if len(parts) > 1:
                    # 清空当前单元格内容
                    for run in paragraph.runs:
                        run.text = ''
                    # 设置第一个部分到当前单元格并添加序号
                    if paragraph.runs:
                        paragraph.paragraph_format.left_indent = 120000
                        paragraph.paragraph_format.first_line_indent = - 120000
                        num_run = paragraph.add_run("1. ")
                        num_run.font.size = Pt(10)  # 使用 Pt 类设置为 10 号字
                        content_run = paragraph.add_run(parts[0])
                        content_run.font.size = Pt(10)  # 使用 Pt 类设置为 10 号字
                    # 为其余部分添加新行并添加序号
                    for i, part in enumerate(parts[1:], start=2):
                        new_row = table.add_row()
                        # 为新行设置 w:cantSplit 属性
                        new_trPr = new_row._element.xpath('.//w:trPr')
                        if not new_trPr:
                            new_trPr = OxmlElement('w:trPr')
                            new_row._element.append(new_trPr)
                        else:
                            new_trPr = new_trPr[0]
                        new_cant_split = new_trPr.xpath('.//w:cantSplit')
                        if not new_cant_split:
                            new_cant_split = OxmlElement('w:cantSplit')
                            new_trPr.append(new_cant_split)
                        new_cell = new_row.cells[col_idx]
                        new_paragraph = new_cell.paragraphs[0]
                        # 动态设置新段落的左缩进和首行缩进
                        digit_count = len(str(i))
                        new_paragraph.paragraph_format.left_indent = 120000 + (digit_count - 1) * 80000
                        new_paragraph.paragraph_format.first_line_indent = - (120000 + (digit_count - 1) * 80000)
                        num_run = new_paragraph.add_run(f"{i}. ")
                        num_run.font.size = Pt(10)  # 使用 Pt 类设置为 10 号字
                        content_run = new_paragraph.add_run(part)
                        content_run.font.size = Pt(10)  # 使用 Pt 类设置为 10 号字
                else:
                    # 若没有分隔符，正常替换
                    replace_paragraph_placeholder(paragraph, placeholder_dict)

# 函数：替换文档中的占位符
def replace_placeholder(doc, placeholder_dict):
    try:
        for paragraph in doc.paragraphs:
            replace_paragraph_placeholder(paragraph, placeholder_dict)
        for table in doc.tables:
            replace_table_placeholder(table, placeholder_dict)
        
        # 检查未替换的占位符
        remaining_placeholders = set()
        pattern = re.compile("|".join(re.escape(placeholder) for placeholder in placeholder_dict.keys()))
        # 检查段落
        for paragraph in doc.paragraphs:
            full_text = ''.join([run.text for run in paragraph.runs])
            for match in pattern.finditer(full_text):
                remaining_placeholders.add(match.group(0))
        # 检查表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = ''.join([run.text for run in paragraph.runs])
                        for match in pattern.finditer(full_text):
                            remaining_placeholders.add(match.group(0))
        
        if remaining_placeholders:
            print(f"未替换的占位符: {remaining_placeholders}")
        else:
            print("所有占位符已成功替换。")

    except Exception as e:
        logging.error(f"替换文档占位符时出错: {e}")

# 函数：获取文档标题
def get_document_title(doc):
    if doc.paragraphs:
        return doc.paragraphs[0].text.strip().replace("/", "_")
    return "Untitled"

# 函数：生成文档
def generate_documents(excel_path, template_path, output_folder):
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel文件未找到: {excel_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Word模板文件未找到: {template_path}")

    df = pd.read_excel(excel_path)
    for index, row in df.iterrows():
        try:
            doc = Document(template_path)
            # 自动根据 Excel 列名生成占位符字典
            placeholder_dict = {f"{{{{{col}}}}}": str(row[col]) for col in df.columns}
            replace_placeholder(doc, placeholder_dict)
            title = get_document_title(doc)
            output_filename = f"{output_folder}/{title}.docx"
            doc.save(output_filename)
            print(f"生成文档：{output_filename}")
        except Exception as e:
            print(f"生成第 {index + 1} 个文档时出错: {e}")

# 函数：获取文档标题
def select_excel_file():
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    if file_path:
        excel_path_entry.delete(0, tk.END)
        excel_path_entry.insert(0, file_path)

def select_template_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        template_path_entry.delete(0, tk.END)
        template_path_entry.insert(0, file_path)

# 新增函数：选择输出文件夹
def select_output_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        output_folder_entry.delete(0, tk.END)
        output_folder_entry.insert(0, folder_path)

def run_generation():
    excel_path = excel_path_entry.get()
    template_path = template_path_entry.get()
    output_folder = output_folder_entry.get()

    # 如果用户没有选择输出文件夹，使用默认路径
    if not output_folder:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_folder = os.path.join(script_dir, 'output')

    # 创建输出文件夹（如果不存在）
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    try:
        print(f"开始处理 {excel_path} 和 {template_path}")
        generate_documents(excel_path, template_path, output_folder)
        print(f"完成处理 {excel_path} 和 {template_path}")
        messagebox.showinfo("完成", "所有文档生成完成。")
    except Exception as e:
        messagebox.showerror("错误", f"处理时出错: {e}")

# 创建主窗口
root = tk.Tk()
root.title("文档生成工具")

# 创建标签和输入框
excel_label = tk.Label(root, text="Excel 文件路径:")
excel_label.pack()
excel_path_entry = tk.Entry(root, width=50)
excel_path_entry.pack()
excel_button = tk.Button(root, text="选择 Excel 文件", command=select_excel_file)
excel_button.pack()

template_label = tk.Label(root, text="Word 模板文件路径:")
template_label.pack()
template_path_entry = tk.Entry(root, width=50)
template_path_entry.pack()
template_button = tk.Button(root, text="选择 Word 模板文件", command=select_template_file)
template_button.pack()

# 新增：输出文件夹标签和输入框
output_folder_label = tk.Label(root, text="输出文件夹路径:")
output_folder_label.pack()
output_folder_entry = tk.Entry(root, width=50)
output_folder_entry.pack()
output_folder_button = tk.Button(root, text="选择输出文件夹", command=select_output_folder)
output_folder_button.pack()

# 创建运行按钮
run_button = tk.Button(root, text="生成文档", command=run_generation)
run_button.pack()

# 运行主循环
root.mainloop()