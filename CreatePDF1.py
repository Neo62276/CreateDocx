import pandas as pd
from docx import Document
import os
import logging
import re
from docx2pdf import convert  # 导入 docx2pdf 库

# 配置日志记录
logging.basicConfig(level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s')

# 函数：替换段落中的占位符
def replace_paragraph_placeholder(paragraph, placeholder_dict):
    try:
        # 构建正则表达式模式
        pattern = re.compile("|".join(re.escape(placeholder) for placeholder in placeholder_dict.keys()))
        # 合并段落中所有 run 的文本
        full_text = ''.join([run.text for run in paragraph.runs])
        # 进行替换
        new_text = pattern.sub(lambda m: placeholder_dict[m.group(0)], full_text)
        # 清空所有 run 的文本
        for run in paragraph.runs:
            run.text = ''
        # 将替换后的文本重新分配到第一个 run 中
        if paragraph.runs:
            paragraph.runs[0].text = new_text
    except Exception as e:
        logging.error(f"替换段落 {paragraph.text[:20]}... 中占位符时出错: {e}")

# 函数：替换表格中的占位符
def replace_table_placeholder(table, placeholder_dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_paragraph_placeholder(paragraph, placeholder_dict)

# 函数：替换文档中的占位符
def replace_placeholder(doc, placeholder_dict):
    try:
        for paragraph in doc.paragraphs:
            replace_paragraph_placeholder(paragraph, placeholder_dict)
        for table in doc.tables:
            replace_table_placeholder(table, placeholder_dict)
        
        # 检查未替换的占位符
        remaining_placeholders = set()
        pattern = re.compile("|".join(re.escape(placeholder) for placeholder in placeholder_dict.keys()))
        # 检查段落
        for paragraph in doc.paragraphs:
            full_text = ''.join([run.text for run in paragraph.runs])
            for match in pattern.finditer(full_text):
                remaining_placeholders.add(match.group(0))
        # 检查表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = ''.join([run.text for run in paragraph.runs])
                        for match in pattern.finditer(full_text):
                            remaining_placeholders.add(match.group(0))
        
        if remaining_placeholders:
            print(f"未替换的占位符: {remaining_placeholders}")
        else:
            print("所有占位符已成功替换。")

    except Exception as e:
        logging.error(f"替换文档占位符时出错: {e}")

# 函数：获取文档标题
def get_document_title(doc):
    if doc.paragraphs:
        return doc.paragraphs[0].text.strip().replace("/", "_")
    return "Untitled"

# 函数：生成文档
def generate_documents(excel_path, template_path, output_folder):
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel文件未找到: {excel_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Word模板文件未找到: {template_path}")

    df = pd.read_excel(excel_path)
    for index, row in df.iterrows():
        try:
            doc = Document(template_path)
            # 自动根据 Excel 列名生成占位符字典
            placeholder_dict = {f"{{{{{col}}}}}": str(row[col]) for col in df.columns}
            replace_placeholder(doc, placeholder_dict)
            title = get_document_title(doc)
            # 保存为临时 Word 文件
            temp_docx_path = f"{output_folder}/{title}.docx"
            doc.save(temp_docx_path)
            # 转换为 PDF 文件
            pdf_path = f"{output_folder}/{title}.pdf"
            convert(temp_docx_path, pdf_path)
            # 删除删除临时 Word 文件的代码
            # os.remove(temp_docx_path)
            print(f"生成文档：{pdf_path}")
        except Exception as e:
            print(f"生成第 {index + 1} 个文档时出错: {e}")

# 主函数
if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_folder = os.path.join(script_dir, 'output')

    # 创建输出文件夹（如果不存在）
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 定义 Excel 文件和模板文件的对应关系
    file_mappings = [
        ('data1.xlsx', 'template1.docx'),
        ('data2.xlsx', 'template2.docx'),
        # 可以继续添加更多对应关系
    ]

    for excel_name, template_name in file_mappings:
        excel_path = os.path.join(script_dir, excel_name)
        template_path = os.path.join(script_dir, template_name)
        try:
            print(f"开始处理 {excel_name} 和 {template_name}")
            generate_documents(excel_path, template_path, output_folder)
            print(f"完成处理 {excel_name} 和 {template_name}")
        except Exception as e:
            print(f"处理 {excel_name} 和 {template_name} 时出错: {e}")